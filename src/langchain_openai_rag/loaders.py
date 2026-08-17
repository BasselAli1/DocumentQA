from __future__ import annotations

from pathlib import Path

import bs4
import requests
from langchain_core.documents import Document


def load_web_page(url: str, *, content_classes: tuple[str, ...] | None = None) -> list[Document]:
    """Load a web page into a LangChain document."""
    bs_kwargs = {}
    if content_classes:
        bs_kwargs["parse_only"] = bs4.SoupStrainer(class_=content_classes)

    response = requests.get(url, timeout=30)
    response.raise_for_status()
    soup = bs4.BeautifulSoup(response.text, "html.parser", **bs_kwargs)
    text = soup.get_text(separator="\n", strip=True)
    return [Document(page_content=text, metadata={"source": url})]


def load_tutorial_source(url: str) -> list[Document]:
    """Load the Lilian Weng article using the same content classes as the tutorial."""
    return load_web_page(
        url,
        content_classes=("post-title", "post-header", "post-content"),
    )


def load_uploaded_document(file_path: Path, *, original_filename: str) -> list[Document]:
    """Load a user-uploaded document into LangChain documents.

    `file_path` is the location of the file on disk (e.g. a temp file), while
    `original_filename` is used to (a) detect the file type from its
    extension and (b) tag the resulting documents with a human-readable
    source name.
    """
    suffix = Path(original_filename).suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path, original_filename)
    if suffix == ".docx":
        return _load_docx(file_path, original_filename)
    if suffix in (".html", ".htm"):
        return _load_html_file(file_path, original_filename)
    if suffix in (".txt", ".md"):
        return _load_text_file(file_path, original_filename)

    raise ValueError(
        f"Unsupported file type '{suffix}'. Supported types: .pdf, .docx, .html, .txt, .md"
    )


def _load_pdf(file_path: Path, source_name: str) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    documents = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": source_name, "page": page_number},
                )
            )
    if not documents:
        raise ValueError(f"No extractable text found in '{source_name}'.")
    return documents


def _load_docx(file_path: Path, source_name: str) -> list[Document]:
    import docx

    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if not text.strip():
        raise ValueError(f"No extractable text found in '{source_name}'.")
    return [Document(page_content=text, metadata={"source": source_name})]


def _load_html_file(file_path: Path, source_name: str) -> list[Document]:
    html = file_path.read_text(encoding="utf-8", errors="ignore")
    soup = bs4.BeautifulSoup(html, "html.parser")
    text = soup.get_text(separator="\n", strip=True)
    if not text.strip():
        raise ValueError(f"No extractable text found in '{source_name}'.")
    return [Document(page_content=text, metadata={"source": source_name})]


def _load_text_file(file_path: Path, source_name: str) -> list[Document]:
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    if not text.strip():
        raise ValueError(f"'{source_name}' is empty.")
    return [Document(page_content=text, metadata={"source": source_name})]
